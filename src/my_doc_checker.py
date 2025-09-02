import requests
import hashlib
from docx import Document
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from reportlab.lib.colors import red, green, black, blue

def download_docx(api_url, save_path):
    response = requests.get(api_url)
    with open(save_path, 'wb') as f:
        f.write(response.content)

def file_hash(file_path):
    sha256 = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            sha256.update(chunk)
    return sha256.hexdigest()


def generate_pdf_report(file1, file2, logs, evidence_pdf):
    c = canvas.Canvas(evidence_pdf, pagesize=A4)
    width, height = A4
    y = height - 2*cm

    # Title
    c.setFont("Helvetica-Bold", 14)
    c.setFillColor(blue)
    c.drawString(2*cm, y, "Document Comparison Report")
    y -= 1*cm

    # Metadata
    c.setFont("Helvetica", 10)
    c.setFillColor(black)
    c.drawString(2*cm, y, f"Timestamp: {datetime.now()}")
    y -= 0.5*cm

    hash_file1 = file_hash(file1)
    hash_file2 = file_hash(file2)

    c.drawString(2*cm, y, f"File 1: {file1}")
    y -= 0.5*cm
    c.drawString(2*cm, y, f"SHA256: {hash_file1}")
    y -= 0.5*cm

    c.drawString(2*cm, y, f"File 2: {file2}")
    y -= 0.5*cm
    c.drawString(2*cm, y, f"SHA256: {hash_file2}")
    y -= 1*cm

    # Results
    c.setFont("Helvetica-Bold", 12)
    c.setFillColor(blue)
    c.drawString(2*cm, y, "Comparison Results:")
    y -= 0.8*cm

    c.setFont("Helvetica", 10)
    for line in logs:
        if y < 2*cm:
            c.showPage()
            y = height - 2*cm
            c.setFont("Helvetica", 10)

        if line.startswith("✅"):
            c.setFillColor(green)
        elif line.startswith("❌"):
            c.setFillColor(red)
        else:
            c.setFillColor(black)

        c.drawString(2*cm, y, line)
        y -= 0.5*cm

    # Final integrity hash
    evidence_hash = hashlib.sha256(("".join(logs) + hash_file1 + hash_file2).encode()).hexdigest()
    c.showPage()
    c.setFont("Helvetica-Bold", 12)
    c.setFillColor(blue)
    c.drawString(2*cm, height - 3*cm, "Integrity Code")
    c.setFont("Helvetica", 10)
    c.setFillColor(black)
    c.drawString(2*cm, height - 4*cm, f"Integrity Code: {evidence_hash}")

    c.save()
    return not any("❌" in log for log in logs), evidence_pdf


def compare_docs(file1, file2, evidence_pdf):
    doc1, doc2 = Document(file1), Document(file2)
    logs = []

    # Paragraphs
    if len(doc1.paragraphs) != len(doc2.paragraphs):
        logs.append(f"❌ Paragraph count mismatch: {len(doc1.paragraphs)} vs {len(doc2.paragraphs)}")

    for idx, (p1, p2) in enumerate(zip(doc1.paragraphs, doc2.paragraphs), start=1):
        if p1.text != p2.text:
            logs.append(f"❌ Paragraph {idx} text mismatch: '{p1.text}' != '{p2.text}'")
        # Bullet / style check
        if p1.style.name != p2.style.name:
            logs.append(f"❌ Paragraph {idx} style mismatch: '{p1.style.name}' vs '{p2.style.name}'")
        for r1, r2 in zip(p1.runs, p2.runs):
            if r1.text != r2.text:
                logs.append(f"❌ Run text mismatch in paragraph {idx}: '{r1.text}' != '{r2.text}'")
            if r1.bold != r2.bold:
                logs.append(f"❌ Bold mismatch in paragraph {idx}: '{r1.text}'")
            if r1.italic != r2.italic:
                logs.append(f"❌ Italic mismatch in paragraph {idx}: '{r1.text}'")
            if r1.underline != r2.underline:
                logs.append(f"❌ Underline mismatch in paragraph {idx}: '{r1.text}'")
            if r1.font.size != r2.font.size:
                logs.append(f"❌ Font size mismatch in paragraph {idx}: '{r1.text}'")
            if r1.font.name != r2.font.name:
                logs.append(f"❌ Font name mismatch in paragraph {idx}: '{r1.text}'")

    # Tables
    if len(doc1.tables) != len(doc2.tables):
        logs.append(f"❌ Table count mismatch: {len(doc1.tables)} vs {len(doc2.tables)}")
    else:
        for t_idx, (t1, t2) in enumerate(zip(doc1.tables, doc2.tables), start=1):
            if len(t1.rows) != len(t2.rows):
                logs.append(f"❌ Table {t_idx} row count mismatch")
            if len(t1.columns) != len(t2.columns):
                logs.append(f"❌ Table {t_idx} column count mismatch")
            for r_idx, (row1, row2) in enumerate(zip(t1.rows, t2.rows), start=1):
                for c_idx, (cell1, cell2) in enumerate(zip(row1.cells, row2.cells), start=1):
                    if cell1.text.strip() != cell2.text.strip():
                        logs.append(f"❌ Table {t_idx} cell({r_idx},{c_idx}) mismatch: '{cell1.text}' vs '{cell2.text}'")

    if not logs:
        logs.append("✅ Documents match perfectly.")

    return generate_pdf_report(file1, file2, logs, evidence_pdf)
